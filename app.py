# Importing necessary libraries
import streamlit as st
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.decomposition import PCA
from sklearn.metrics import silhouette_score
from fpdf import FPDF
import io
from docx import Document

# Set the Streamlit page configuration
st.set_page_config(
    page_title="Customer Segmentation App",
    page_icon=":chart_with_upwards_trend:",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Add a custom logo to the sidebar
st.sidebar.image("logo.png", use_container_width=True)
st.sidebar.title("Customer Segmentation App")

# Provide a guide for the file format and contents
st.sidebar.subheader("File Requirements")
st.sidebar.write("Please upload a dataset with the following columns and data types:")
file_guide = pd.DataFrame({
    "Column Name": ['Education', 'Income', 'Recency', 'Complain', 'Response', 'Age', 'Spent', 'Living_With', 'Children', 'Family_Size', 'Is_Parent'],
    "Data Type": ['Categorical', 'Numeric', 'Numeric', 'Binary (0/1)', 'Binary (0/1)', 'Numeric', 'Numeric', 'Categorical', 'Numeric', 'Numeric', 'Binary (0/1)']
})
st.sidebar.dataframe(file_guide)

# Upload dataset
st.sidebar.subheader("Upload Dataset")
uploaded_file = st.sidebar.file_uploader("Choose a file", type=["csv", "xlsx"])

# Load dataset
@st.cache_data
def load_data(file):
    if file.name.endswith('.csv'):
        return pd.read_csv(file)
    elif file.name.endswith('.xlsx'):
        return pd.read_excel(file)

if uploaded_file:
    df = load_data(uploaded_file)
    st.write("### Uploaded Dataset Preview")
    st.dataframe(df.head())

    # Validate dataset contents
    missing_columns = [col for col in file_guide['Column Name'] if col not in df.columns]
    if missing_columns:
        st.error(f"The following required columns are missing: {', '.join(missing_columns)}")
    else:
        st.success("All required columns are present!")

        # Sidebar: Select features for clustering
        st.sidebar.subheader("Select Features for Clustering")
        features = st.sidebar.multiselect(
            "Choose features to include in the clustering:",
            options=df.columns,
            default=['Education', 'Income', 'Recency', 'Complain', 'Response', 'Age', 'Spent', 'Living_With', 'Children', 'Family_Size', 'Is_Parent']
        )

        if features:
            st.write(f"### Selected Features: {', '.join(features)}")
            
            # Data Preprocessing
            imputer = SimpleImputer(strategy='mean')
            scaler = StandardScaler()

            data = imputer.fit_transform(df[features])
            scaled_data = scaler.fit_transform(data)

            # KMeans Clustering
            st.sidebar.subheader("KMeans Clustering Parameters")
            n_clusters = st.sidebar.slider("Number of Clusters", 2, 10, value=3)
            kmeans = KMeans(n_clusters=n_clusters, random_state=42)

            if st.button("Run Clustering"):
                cluster_labels = kmeans.fit_predict(scaled_data)
                df['Cluster'] = cluster_labels

                # Display cluster counts
                st.write("### Cluster Distribution")
                st.bar_chart(df['Cluster'].value_counts())

                # Visualize Clusters using PCA
                st.write("### PCA Cluster Visualization")
                pca = PCA(n_components=2)
                pca_data = pca.fit_transform(scaled_data)
                pca_df = pd.DataFrame(pca_data, columns=['PCA1', 'PCA2'])
                pca_df['Cluster'] = cluster_labels

                fig = px.scatter(
                    pca_df, x='PCA1', y='PCA2',
                    color=pca_df['Cluster'].astype(str),
                    title="Clusters Visualized in 2D Space (PCA)",
                    template='plotly_dark',
                    labels={"Cluster": "Cluster"},
                    width=800, height=500
                )
                st.plotly_chart(fig)

                # Predict Button
                if st.button("Predict for New Data"):
                    st.write("### Prediction Functionality Coming Soon!")

                # Dashboard Visualization
                st.write("### Dashboard")
                cluster_summary = df.groupby('Cluster').mean()
                st.write("**Cluster Summary:**")
                st.dataframe(cluster_summary)

                # Save results
                st.write("### Save Results")
                buffer = io.BytesIO()

                save_format = st.radio("Select format for saving:", ['Excel', 'CSV', 'PDF', 'DOC'])

                if st.button("Save Results"):
                    if save_format == 'Excel':
                        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                            df.to_excel(writer, sheet_name='Clusters', index=False)
                            writer.save()
                        st.download_button(
                            label="Download Excel File",
                            data=buffer.getvalue(),
                            file_name="clustered_data.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                        )
                    elif save_format == 'CSV':
                        csv_data = df.to_csv(index=False)
                        st.download_button(
                            label="Download CSV File",
                            data=csv_data,
                            file_name="clustered_data.csv",
                            mime="text/csv"
                        )
                    elif save_format == 'PDF':
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_font("Arial", size=12)
                        for row in df.itertuples():
                            pdf.cell(200, 10, txt=str(row), ln=True)
                        pdf_output = buffer
                        pdf.output(pdf_output)
                        st.download_button(
                            label="Download PDF File",
                            data=buffer.getvalue(),
                            file_name="clustered_data.pdf",
                            mime="application/pdf"
                        )
                    elif save_format == 'DOC':
                        doc = Document()
                        doc.add_heading("Clustered Data", level=1)
                        for row in df.itertuples():
                            doc.add_paragraph(str(row))
                        doc.save(buffer)
                        st.download_button(
                            label="Download DOC File",
                            data=buffer.getvalue(),
                            file_name="clustered_data.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
else:
    st.write("## Please upload a dataset to proceed.")
